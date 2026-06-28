"""Generate individual video-submission documents as Word (.docx) files.

Editable format expected by the submission portal: project/member details,
the cloud-drive video link, and short Output / Novelty / Contribution sections.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_TITLE = "EEG-Based Emotion Recognition Using Machine Learning and Deep Learning"
GROUP = "Cognitive Squad"
DEGREE = "MSc in Information Technology"
INSTITUTE = "Sri Lanka Institute of Information Technology (SLIIT)"
SUBJECT = "Artificial Intelligence"
REPO = "https://github.com/Malinga5194/cognitive-squad-eeg-emotion-recognition"

NAVY = RGBColor(0x14, 0x1E, 0x61)
BLUE = RGBColor(0x15, 0x65, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55)

members = [
    {
        "name": "D.S.M. Perera",
        "id": "MS26906294",
        "role": "Team Lead / System Architect",
        "docx": "paper/submission_Perera.docx",
        "link": "https://mysliit-my.sharepoint.com/:v:/g/personal/ms26906294_my_sliit_lk/IQDobFgEMcFzQaOuuj6aOwTyAQYkcA133_cnh38444io-cw",
        "output": (
            "The project is an EEG-Based Emotion Recognition system that classifies three "
            "emotional states - Positive, Neutral, and Negative - directly from brainwave "
            "signals recorded by a consumer EEG headband. The system compares seven AI models "
            "through an interactive web dashboard and achieves 98.83% classification accuracy."
        ),
        "novelty": (
            "The novelty is the comprehensive comparison of seven models, ranging from "
            "traditional SVM and Random Forest to a state-of-the-art Transformer with "
            "self-attention. A key finding is that a traditional Random Forest matches the "
            "Transformer at 98.83%, showing that feature quality matters more than model complexity."
        ),
        "contribution": (
            "I designed the complete system architecture, deciding the project structure and "
            "frameworks and organising the code into modular components. I implemented three deep "
            "learning models in PyTorch - a Transformer with multi-head self-attention, a hybrid "
            "CNN-LSTM, and a bidirectional LSTM - and built the interactive Streamlit web dashboard "
            "and managed the team GitHub repository."
        ),
    },
]


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1565C0")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")  # 9pt
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def detail_row(doc, label, value, bold_value=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = GRAY
    r2 = p.add_run(value)
    r2.bold = bold_value
    r2.font.size = Pt(10.5)
    return p


def section(doc, heading, body):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    rh = h.add_run(heading)
    rh.bold = True
    rh.font.size = Pt(12.5)
    rh.font.color.rgb = NAVY
    b = doc.add_paragraph(body)
    b.paragraph_format.space_after = Pt(4)
    for run in b.runs:
        run.font.size = Pt(10.5)
    b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


for m in members:
    doc = Document()
    section_obj = doc.sections[0]
    section_obj.top_margin = Inches(0.7)
    section_obj.bottom_margin = Inches(0.7)
    section_obj.left_margin = Inches(0.9)
    section_obj.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ---- Header band (single-cell shaded table) ----
    header = doc.add_table(rows=1, cols=1)
    cell = header.rows[0].cells[0]
    set_cell_bg(cell, "141E61")
    p = cell.paragraphs[0]
    r = p.add_run("Video Submission")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(f"{GROUP}  |  {DEGREE}  |  SLIIT")
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0xD5, 0xDA, 0xF0)

    doc.add_paragraph()

    # ---- Details ----
    detail_row(doc, "Project Title:", PROJECT_TITLE, bold_value=True)
    detail_row(doc, "Subject:", SUBJECT)
    detail_row(doc, "Group:", GROUP)
    detail_row(doc, "Institute:", INSTITUTE)
    doc.add_paragraph()
    detail_row(doc, "Member Name:", m["name"], bold_value=True)
    detail_row(doc, "Student ID:", m["id"])
    detail_row(doc, "Role:", m["role"])
    detail_row(doc, "Video Duration:", "3 minutes")

    # ---- Video link box ----
    doc.add_paragraph()
    link_tbl = doc.add_table(rows=1, cols=1)
    link_tbl.style = "Table Grid"
    lcell = link_tbl.rows[0].cells[0]
    set_cell_bg(lcell, "EEF3FB")
    lp = lcell.paragraphs[0]
    lr = lp.add_run("Video Link (OneDrive):")
    lr.bold = True
    lr.font.size = Pt(11)
    lr.font.color.rgb = BLUE
    link_para = lcell.add_paragraph()
    add_hyperlink(link_para, m["link"], m["link"])

    # ---- Sections ----
    section(doc, "1. Project Output", m["output"])
    section(doc, "2. Novelty", m["novelty"])
    section(doc, "3. Individual Contribution", m["contribution"])

    # ---- Footer note ----
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fr = fp.add_run("Project repository: " + REPO)
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = GRAY

    doc.save(m["docx"])
    print(f"Created: {m['docx']}")

print("\nWord submission document generated!")
